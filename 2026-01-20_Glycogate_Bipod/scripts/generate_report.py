#!/usr/bin/env python3
"""
Bipod 10ns MD 시뮬레이션 결과 보고서 생성 (Word 파일)
"""
import os
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 파일 경로
RESULTS_DIR = "/home/pjho3tr/projects/Drug/2026-01-20_Glycogate_Bipod/results/md_bipod_10ns"
ARM_CSV = os.path.join(RESULTS_DIR, "bipod_arm_analysis.csv")
GLUCOSE_CSV = os.path.join(RESULTS_DIR, "glucose_distance.csv")
PLOT1 = os.path.join(RESULTS_DIR, "bipod_analysis_plots.png")
PLOT2 = os.path.join(RESULTS_DIR, "bipod_analysis_distributions.png")
OUTPUT_DOCX = os.path.join(RESULTS_DIR, "Bipod_10ns_MD_Report.docx")

def add_heading_with_style(doc, text, level=1):
    """스타일이 적용된 제목 추가"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_table_with_data(doc, data, headers):
    """데이터 테이블 추가"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # 헤더 추가
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # 헤더 볼드 처리
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # 데이터 추가
    for row_data in data:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)
    
    return table

def main():
    # 데이터 로드
    arm_data = pd.read_csv(ARM_CSV)
    glucose_data = pd.read_csv(GLUCOSE_CSV)
    
    # Word 문서 생성
    doc = Document()
    
    # 제목
    title = doc.add_heading('Bipod MD Simulation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 부제목
    subtitle = doc.add_paragraph('10 ns Molecular Dynamics Simulation of TRIS-PEG24-L-glucose Bipod')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(14)
    subtitle_format.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # 빈 줄
    
    # 1. Executive Summary
    add_heading_with_style(doc, '1. Executive Summary', level=1)
    
    summary_text = f"""
This report presents the results of a 10 ns molecular dynamics (MD) simulation of a bipod structure 
consisting of TRIS core with two PEG24-L-glucose arms. The simulation was performed using OpenMM 
with CHARMM36 force field in a 120 Å cubic water box with 150 mM NaCl.

Key findings:
• Both arms extended well with average lengths of ~78 Å (7.8 nm)
• Distance between two L-glucose groups: {glucose_data['glucose_distance_A'].mean():.1f} ± {glucose_data['glucose_distance_A'].std():.1f} Å
• Radius of gyration: {arm_data['Rg_A'].mean():.1f} ± {arm_data['Rg_A'].std():.1f} Å
• Both arms showed 100% probability of extending beyond 50 Å (5 nm)
"""
    doc.add_paragraph(summary_text.strip())
    
    # 2. System Setup
    add_heading_with_style(doc, '2. System Setup', level=1)
    
    setup_data = [
        ['Parameter', 'Value'],
        ['Molecule', 'TRIS-PEG24-L-glucose Bipod'],
        ['Force Field', 'CHARMM36'],
        ['Water Model', 'TIP3P'],
        ['Box Size', '120 × 120 × 120 Å³'],
        ['Salt Concentration', '150 mM NaCl'],
        ['Temperature', '300 K'],
        ['Pressure', '1 atm'],
        ['Simulation Length', '10 ns'],
        ['Time Step', '2 fs'],
        ['Total Frames', '100']
    ]
    
    p = doc.add_paragraph('Simulation parameters:')
    table = add_table_with_data(doc, setup_data[1:], setup_data[0])
    doc.add_paragraph()
    
    # 3. Results
    add_heading_with_style(doc, '3. Results', level=1)
    
    # 3.1 Arm Length Analysis
    add_heading_with_style(doc, '3.1 Individual Arm Lengths', level=2)
    
    arm1_mean = arm_data['arm1_A'].mean()
    arm1_std = arm_data['arm1_A'].std()
    arm1_min = arm_data['arm1_A'].min()
    arm1_max = arm_data['arm1_A'].max()
    
    arm2_mean = arm_data['arm2_A'].mean()
    arm2_std = arm_data['arm2_A'].std()
    arm2_min = arm_data['arm2_A'].min()
    arm2_max = arm_data['arm2_A'].max()
    
    arm_stats = [
        ['Metric', 'Arm 1 (Å)', 'Arm 2 (Å)'],
        ['Mean ± Std', f'{arm1_mean:.2f} ± {arm1_std:.2f}', f'{arm2_mean:.2f} ± {arm2_std:.2f}'],
        ['Minimum', f'{arm1_min:.2f}', f'{arm2_min:.2f}'],
        ['Maximum', f'{arm1_max:.2f}', f'{arm2_max:.2f}'],
        ['P(≥ 40 Å)', '100%', '100%'],
        ['P(≥ 50 Å)', '100%', '100%']
    ]
    
    doc.add_paragraph('Statistical summary of arm lengths:')
    add_table_with_data(doc, arm_stats[1:], arm_stats[0])
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Key observations:').bold = True
    doc.add_paragraph('• Both arms showed similar average lengths (~78 Å)', style='List Bullet')
    doc.add_paragraph('• Arms exhibited dynamic behavior, fluctuating between 50-109 Å', style='List Bullet')
    doc.add_paragraph('• 100% of conformations had both arms extended beyond 50 Å', style='List Bullet')
    
    doc.add_paragraph()
    
    # 3.2 Glucose Distance
    add_heading_with_style(doc, '3.2 Distance Between L-glucose Groups', level=2)
    
    gluc_mean = glucose_data['glucose_distance_A'].mean()
    gluc_std = glucose_data['glucose_distance_A'].std()
    gluc_min = glucose_data['glucose_distance_A'].min()
    gluc_max = glucose_data['glucose_distance_A'].max()
    
    glucose_stats = [
        ['Metric', 'Value (Å)'],
        ['Mean ± Std', f'{gluc_mean:.2f} ± {gluc_std:.2f}'],
        ['Minimum', f'{gluc_min:.2f}'],
        ['Maximum', f'{gluc_max:.2f}'],
        ['P(100-150 Å)', '99%'],
        ['P(≥ 150 Å)', '1%']
    ]
    
    doc.add_paragraph('Statistical summary of glucose-glucose distance:')
    add_table_with_data(doc, glucose_stats[1:], glucose_stats[0])
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Key observations:').bold = True
    doc.add_paragraph(f'• Average distance: {gluc_mean:.1f} Å (≈ {gluc_mean/10:.1f} nm)', style='List Bullet')
    doc.add_paragraph('• Two glucose groups remained well-separated throughout simulation', style='List Bullet')
    doc.add_paragraph('• Distance varied between 107-170 Å, showing dynamic behavior', style='List Bullet')
    
    doc.add_paragraph()
    
    # 3.3 Radius of Gyration
    add_heading_with_style(doc, '3.3 Radius of Gyration', level=2)
    
    rg_mean = arm_data['Rg_A'].mean()
    rg_std = arm_data['Rg_A'].std()
    rg_min = arm_data['Rg_A'].min()
    rg_max = arm_data['Rg_A'].max()
    
    rg_stats = [
        ['Metric', 'Value (Å)'],
        ['Mean ± Std', f'{rg_mean:.2f} ± {rg_std:.2f}'],
        ['Minimum', f'{rg_min:.2f}'],
        ['Maximum', f'{rg_max:.2f}']
    ]
    
    doc.add_paragraph('Statistical summary of radius of gyration:')
    add_table_with_data(doc, rg_stats[1:], rg_stats[0])
    
    doc.add_paragraph()
    doc.add_paragraph(
        f'The radius of gyration remained relatively stable at {rg_mean:.1f} ± {rg_std:.1f} Å, '
        'indicating a consistent overall size of the bipod structure during the simulation.'
    )
    
    # 4. Figures
    doc.add_page_break()
    add_heading_with_style(doc, '4. Figures', level=1)
    
    # Figure 1
    add_heading_with_style(doc, 'Figure 1. Time Series Analysis', level=2)
    if os.path.exists(PLOT1):
        doc.add_picture(PLOT1, width=Inches(6.5))
        doc.add_paragraph(
            'Time evolution of (A) individual arm lengths, (B) maximum arm length, '
            '(C) distance between two L-glucose groups, and (D) radius of gyration over 10 ns simulation.',
            style='Caption'
        )
    
    doc.add_page_break()
    
    # Figure 2
    add_heading_with_style(doc, 'Figure 2. Distribution Analysis', level=2)
    if os.path.exists(PLOT2):
        doc.add_picture(PLOT2, width=Inches(6.5))
        doc.add_paragraph(
            'Distribution histograms of (A) Arm 1 length, (B) Arm 2 length, '
            '(C) glucose-glucose distance, and (D) radius of gyration.',
            style='Caption'
        )
    
    # 5. Conclusions
    doc.add_page_break()
    add_heading_with_style(doc, '5. Conclusions', level=1)
    
    conclusions = """
The 10 ns MD simulation of the TRIS-PEG24-L-glucose bipod structure revealed the following key findings:

1. Structural Extension:
   • Both arms extended well with average lengths of approximately 78 Å (7.8 nm)
   • Arms showed 100% probability of extending beyond 50 Å (5 nm)
   • Dynamic behavior observed with arm lengths fluctuating between 50-109 Å

2. Glucose Separation:
   • Two L-glucose groups maintained an average separation of 125 Å (12.5 nm)
   • Distance remained predominantly in the 100-150 Å range (99% of frames)
   • Maximum separation reached 170 Å, indicating high flexibility

3. Overall Structure:
   • Radius of gyration: 52.7 ± 3.0 Å
   • Bipod structure maintained extended conformation throughout simulation
   • Both arms showed independent dynamic behavior

4. Recommendations:
   • Extend simulation to 100-200 ns for better statistical sampling
   • Compare with single-arm (1-arm) structure to assess bipod-specific effects
   • Consider analyzing arm-arm correlation and potential interactions
"""
    
    doc.add_paragraph(conclusions.strip())
    
    # 6. Technical Notes
    doc.add_page_break()
    add_heading_with_style(doc, '6. Technical Notes', level=1)
    
    tech_notes = """
Energy Minimization:
• Initial energy: 2.42 × 10²⁷ kJ/mol (before box correction)
• Final energy after minimization: -2,860,925 kJ/mol
• Three-stage minimization protocol applied

CGenFF Parameters:
• Penalty score: 149.5 (high - requires validation)
• Recommendation: Parameter optimization may improve accuracy

Simulation Protocol:
• Platform: CUDA (GPU acceleration)
• Integrator: Langevin (300 K)
• Barostat: Monte Carlo (1 atm)
• Nonbonded method: PME
• Cutoff: 12 Å
• Constraints: H-bonds
"""
    
    doc.add_paragraph(tech_notes.strip())
    
    # 저장
    doc.save(OUTPUT_DOCX)
    print(f"\n✅ Report saved: {OUTPUT_DOCX}")
    print(f"   File size: {os.path.getsize(OUTPUT_DOCX) / 1024:.1f} KB")

if __name__ == "__main__":
    main()
