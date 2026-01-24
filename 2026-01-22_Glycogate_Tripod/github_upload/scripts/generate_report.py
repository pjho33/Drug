#!/usr/bin/env python3
"""
Tripod MD 시뮬레이션 결과 워드 보고서 생성
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = "/home/pjho3tr/projects/Drug/2026-01-22_Glycogate_Tripod/results"
OUTPUT_DOCX = f"{BASE_DIR}/Tripod_MD_Simulation_Report.docx"

# 데이터 로드
data_1ns = pd.read_csv(f"{BASE_DIR}/md_tripod_1ns/tripod_arm_analysis.csv")
data_10ns = pd.read_csv(f"{BASE_DIR}/md_tripod_10ns/tripod_arm_analysis.csv")
data_100ns = pd.read_csv(f"{BASE_DIR}/md_tripod_100ns/tripod_arm_analysis.csv")

glucose_1ns = pd.read_csv(f"{BASE_DIR}/md_tripod_1ns/glucose_distances.csv")
glucose_10ns = pd.read_csv(f"{BASE_DIR}/md_tripod_10ns/glucose_distances.csv")
glucose_100ns = pd.read_csv(f"{BASE_DIR}/md_tripod_100ns/glucose_distances.csv")

def add_heading(doc, text, level=1):
    """제목 추가"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False):
    """문단 추가"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_table_from_data(doc, headers, rows):
    """테이블 추가"""
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # 헤더
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        run = cell.paragraphs[0].runs[0]
        run.bold = True
    
    # 데이터
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.rows[i+1].cells[j].text = str(value)
    
    return table

def main():
    doc = Document()
    
    # 제목
    title = doc.add_heading('Tripod MD Simulation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 부제목
    subtitle = doc.add_paragraph('Molecular Dynamics Simulation Analysis')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    date_p = doc.add_paragraph('Date: January 24, 2026')
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # 1. Executive Summary
    add_heading(doc, '1. Executive Summary', 1)
    add_paragraph(doc, 
        'This report presents the results of molecular dynamics (MD) simulations '
        'performed on a Tripod structure at three different timescales: 1ns, 10ns, and 100ns. '
        'The simulations were conducted using OpenMM with CUDA acceleration on GPU.')
    
    doc.add_paragraph()
    add_paragraph(doc, 'Key Findings:', bold=True)
    
    findings = doc.add_paragraph(style='List Bullet')
    findings.add_run('The Tripod structure shows significant time-dependent conformational changes')
    
    findings = doc.add_paragraph(style='List Bullet')
    findings.add_run('Initial 1ns simulation shows extended conformation (74.3 Å)')
    
    findings = doc.add_paragraph(style='List Bullet')
    findings.add_run('100ns simulation reveals highly compact equilibrium structure (30.4 Å)')
    
    findings = doc.add_paragraph(style='List Bullet')
    findings.add_run('Only 5.4% of conformations exceed 50 Å in 100ns simulation')
    
    doc.add_page_break()
    
    # 2. Simulation Details
    add_heading(doc, '2. Simulation Details', 1)
    
    add_heading(doc, '2.1 System Setup', 2)
    add_paragraph(doc, 'Force Field: CHARMM36')
    add_paragraph(doc, 'Water Model: TIP3P')
    add_paragraph(doc, 'Box Size: 120 × 120 × 120 Å³')
    add_paragraph(doc, 'Temperature: 303.15 K')
    add_paragraph(doc, 'Pressure: 1.0 bar')
    add_paragraph(doc, 'Platform: CUDA (GPU acceleration)')
    
    doc.add_paragraph()
    add_heading(doc, '2.2 Simulation Parameters', 2)
    
    sim_data = [
        ['Simulation', 'Total Steps', 'Time Step', 'Total Time', 'Frames'],
        ['1ns', '250,000', '4 fs', '1 ns', '10'],
        ['10ns', '2,500,000', '4 fs', '10 ns', '100'],
        ['100ns', '25,000,000', '4 fs', '100 ns', '1,000']
    ]
    add_table_from_data(doc, sim_data[0], sim_data[1:])
    
    doc.add_page_break()
    
    # 3. Results
    add_heading(doc, '3. Results', 1)
    
    add_heading(doc, '3.1 Arm Length Analysis', 2)
    
    # 통계 계산
    stats_1ns = {
        'max_mean': data_1ns['max_arm_A'].mean(),
        'max_std': data_1ns['max_arm_A'].std(),
        'max_min': data_1ns['max_arm_A'].min(),
        'max_max': data_1ns['max_arm_A'].max(),
        'p40': (data_1ns['max_arm_A'] >= 40).mean() * 100,
        'p50': (data_1ns['max_arm_A'] >= 50).mean() * 100
    }
    
    stats_10ns = {
        'max_mean': data_10ns['max_arm_A'].mean(),
        'max_std': data_10ns['max_arm_A'].std(),
        'max_min': data_10ns['max_arm_A'].min(),
        'max_max': data_10ns['max_arm_A'].max(),
        'p40': (data_10ns['max_arm_A'] >= 40).mean() * 100,
        'p50': (data_10ns['max_arm_A'] >= 50).mean() * 100
    }
    
    stats_100ns = {
        'max_mean': data_100ns['max_arm_A'].mean(),
        'max_std': data_100ns['max_arm_A'].std(),
        'max_min': data_100ns['max_arm_A'].min(),
        'max_max': data_100ns['max_arm_A'].max(),
        'p40': (data_100ns['max_arm_A'] >= 40).mean() * 100,
        'p50': (data_100ns['max_arm_A'] >= 50).mean() * 100
    }
    
    arm_data = [
        ['Metric', '1ns', '10ns', '100ns'],
        ['Mean ± Std (Å)', 
         f'{stats_1ns["max_mean"]:.1f} ± {stats_1ns["max_std"]:.1f}',
         f'{stats_10ns["max_mean"]:.1f} ± {stats_10ns["max_std"]:.1f}',
         f'{stats_100ns["max_mean"]:.1f} ± {stats_100ns["max_std"]:.1f}'],
        ['Min - Max (Å)', 
         f'{stats_1ns["max_min"]:.1f} - {stats_1ns["max_max"]:.1f}',
         f'{stats_10ns["max_min"]:.1f} - {stats_10ns["max_max"]:.1f}',
         f'{stats_100ns["max_min"]:.1f} - {stats_100ns["max_max"]:.1f}'],
        ['P(≥40Å) %', 
         f'{stats_1ns["p40"]:.1f}',
         f'{stats_10ns["p40"]:.1f}',
         f'{stats_100ns["p40"]:.1f}'],
        ['P(≥50Å) %', 
         f'{stats_1ns["p50"]:.1f}',
         f'{stats_10ns["p50"]:.1f}',
         f'{stats_100ns["p50"]:.1f}']
    ]
    add_table_from_data(doc, arm_data[0], arm_data[1:])
    
    doc.add_paragraph()
    add_heading(doc, '3.2 Radius of Gyration', 2)
    
    rg_data = [
        ['Simulation', 'Mean ± Std (Å)', 'Min - Max (Å)'],
        ['1ns', 
         f'{data_1ns["Rg_A"].mean():.1f} ± {data_1ns["Rg_A"].std():.1f}',
         f'{data_1ns["Rg_A"].min():.1f} - {data_1ns["Rg_A"].max():.1f}'],
        ['10ns', 
         f'{data_10ns["Rg_A"].mean():.1f} ± {data_10ns["Rg_A"].std():.1f}',
         f'{data_10ns["Rg_A"].min():.1f} - {data_10ns["Rg_A"].max():.1f}'],
        ['100ns', 
         f'{data_100ns["Rg_A"].mean():.1f} ± {data_100ns["Rg_A"].std():.1f}',
         f'{data_100ns["Rg_A"].min():.1f} - {data_100ns["Rg_A"].max():.1f}']
    ]
    add_table_from_data(doc, rg_data[0], rg_data[1:])
    
    doc.add_paragraph()
    add_heading(doc, '3.3 Glucose-Glucose Distances', 2)
    
    gluc_data = [
        ['Distance', '1ns (Å)', '10ns (Å)', '100ns (Å)'],
        ['G1-G2', 
         f'{glucose_1ns["glucose12_A"].mean():.1f} ± {glucose_1ns["glucose12_A"].std():.1f}',
         f'{glucose_10ns["glucose12_A"].mean():.1f} ± {glucose_10ns["glucose12_A"].std():.1f}',
         f'{glucose_100ns["glucose12_A"].mean():.1f} ± {glucose_100ns["glucose12_A"].std():.1f}'],
        ['G1-G3', 
         f'{glucose_1ns["glucose13_A"].mean():.1f} ± {glucose_1ns["glucose13_A"].std():.1f}',
         f'{glucose_10ns["glucose13_A"].mean():.1f} ± {glucose_10ns["glucose13_A"].std():.1f}',
         f'{glucose_100ns["glucose13_A"].mean():.1f} ± {glucose_100ns["glucose13_A"].std():.1f}'],
        ['G2-G3', 
         f'{glucose_1ns["glucose23_A"].mean():.1f} ± {glucose_1ns["glucose23_A"].std():.1f}',
         f'{glucose_10ns["glucose23_A"].mean():.1f} ± {glucose_10ns["glucose23_A"].std():.1f}',
         f'{glucose_100ns["glucose23_A"].mean():.1f} ± {glucose_100ns["glucose23_A"].std():.1f}']
    ]
    add_table_from_data(doc, gluc_data[0], gluc_data[1:])
    
    doc.add_page_break()
    
    # 4. Discussion
    add_heading(doc, '4. Discussion', 1)
    
    add_heading(doc, '4.1 Time-Dependent Conformational Changes', 2)
    add_paragraph(doc, 
        'The Tripod structure exhibits significant conformational changes over the simulation timescale. '
        'The initial 1ns simulation shows an extended conformation with a maximum arm length of 74.3 Å, '
        'which is likely a non-equilibrium state immediately following energy minimization and equilibration.')
    
    doc.add_paragraph()
    add_paragraph(doc, 
        'As the simulation progresses to 10ns and 100ns, the structure undergoes substantial compaction. '
        'The 100ns simulation reveals a highly compact equilibrium structure with a mean maximum arm length '
        'of only 30.4 Å, representing a 59% reduction from the 1ns value.')
    
    doc.add_paragraph()
    add_heading(doc, '4.2 Structural Compactness', 2)
    add_paragraph(doc, 
        'The radius of gyration analysis confirms the trend toward compactness. The Rg decreases from '
        '44.1 Å (1ns) to 20.5 Å (100ns), indicating that the Tripod structure strongly prefers a compact '
        'conformation in aqueous solution.')
    
    doc.add_paragraph()
    add_heading(doc, '4.3 Extension Probability', 2)
    add_paragraph(doc, 
        'The probability of finding extended conformations (≥50 Å) decreases dramatically from 100% in '
        'the 1ns simulation to only 5.4% in the 100ns simulation. This suggests that the Tripod structure '
        'spends most of its time in a collapsed state.')
    
    doc.add_page_break()
    
    # 5. Conclusions
    add_heading(doc, '5. Conclusions', 1)
    
    conclusions = doc.add_paragraph(style='List Number')
    conclusions.add_run('The Tripod structure exhibits strong preference for compact conformations in aqueous solution.')
    
    conclusions = doc.add_paragraph(style='List Number')
    conclusions.add_run('Extended conformations observed in short simulations (1ns) are non-equilibrium states.')
    
    conclusions = doc.add_paragraph(style='List Number')
    conclusions.add_run('Long-timescale simulations (100ns) are necessary to capture equilibrium behavior.')
    
    conclusions = doc.add_paragraph(style='List Number')
    conclusions.add_run('The three arms of the Tripod structure show symmetric behavior with similar lengths.')
    
    conclusions = doc.add_paragraph(style='List Number')
    conclusions.add_run('Compared to Bipod structures, Tripod shows significantly less extension, likely due to steric constraints from the third arm.')
    
    doc.add_page_break()
    
    # 6. Figures
    add_heading(doc, '6. Figures', 1)
    
    # 비교 플롯 추가
    comparison_plot = f"{BASE_DIR}/tripod_comparison_all.png"
    if os.path.exists(comparison_plot):
        add_heading(doc, 'Figure 1: Comparison of 1ns, 10ns, and 100ns Simulations', 2)
        doc.add_picture(comparison_plot, width=Inches(6.5))
        doc.add_paragraph()
    
    # 개별 플롯들
    for sim_name in ['1ns', '10ns', '100ns']:
        plot_file = f"{BASE_DIR}/md_tripod_{sim_name}/tripod_{sim_name}_plots.png"
        if os.path.exists(plot_file):
            add_heading(doc, f'Figure: {sim_name.upper()} Simulation Detailed Analysis', 2)
            doc.add_picture(plot_file, width=Inches(6.5))
            doc.add_paragraph()
            if sim_name != '100ns':
                doc.add_page_break()
    
    # 저장
    doc.save(OUTPUT_DOCX)
    print(f"\n✅ Word report saved: {OUTPUT_DOCX}")

if __name__ == "__main__":
    main()
